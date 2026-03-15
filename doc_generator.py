# doc_generator.py
from docx import Document
from docx.shared import Pt
import os
import uuid
import re

GENERATED_DIR = "generated"
os.makedirs(GENERATED_DIR, exist_ok=True)

def sanitize_filename(name):
    return re.sub(r'[^a-zA-Z0-9_-]', '_', name)

def generate_doc(topic, lesson_text):
    """
    Generate a professional Word (.docx) lesson plan from AI text.
    - Structured headings (Objectives, Activities, Assessment)
    - Readable 21st-century style
    - Returns the path to the saved document
    """

    doc = Document()
    doc.add_heading(f"Lesson Plan: {topic}", level=0)

    # Split lesson text into sections using double newlines
    sections = [s.strip() for s in lesson_text.split("\n\n") if s.strip()]

    for section in sections:
        lines = section.split("\n")
        # First line treated as section heading
        heading = lines[0][:60]  # truncate very long headings
        doc.add_heading(heading, level=1)
        # Remaining lines as content
        for line in lines[1:]:
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = Pt(14)

    # Safe filename
    file_name = f"{sanitize_filename(topic)}_{uuid.uuid4().hex[:5]}.docx"
    path = os.path.join(GENERATED_DIR, file_name)
    doc.save(path)
    return path